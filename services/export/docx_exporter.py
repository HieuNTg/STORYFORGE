"""Export story as DOCX (Word) using python-docx. Vietnamese-friendly out of the box."""

import logging
import os
from typing import Union
from models.schemas import StoryDraft, EnhancedStory, Character

logger = logging.getLogger(__name__)


class DOCXExporter:
    """Export story as a Word document with chapter headings + paragraphs."""

    @staticmethod
    def export(
        story: Union[StoryDraft, EnhancedStory],
        output_path: str,
        characters: list[Character] = None,
        author: str = "StoryForge",
    ) -> str:
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Run: pip install python-docx")
            return ""

        doc = Document()

        # Base style — Calibri renders Vietnamese diacritics correctly on Win/Mac.
        try:
            normal = doc.styles["Normal"]
            normal.font.name = "Calibri"
            normal.font.size = Pt(12)
        except Exception:
            pass

        # Title page
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.add_run(story.title or "Untitled")
        title_run.bold = True
        title_run.font.size = Pt(28)

        if hasattr(story, "genre") and story.genre:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub_run = sub.add_run(f"Thể loại: {story.genre}")
            sub_run.italic = True
            sub_run.font.size = Pt(13)

        author_para = doc.add_paragraph()
        author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_run = author_para.add_run(f"Tác giả: {author}")
        author_run.font.size = Pt(12)

        if hasattr(story, "synopsis") and story.synopsis:
            doc.add_paragraph()
            doc.add_heading("Tóm tắt", level=2)
            doc.add_paragraph(story.synopsis)

        doc.add_page_break()

        # Characters
        if characters:
            doc.add_heading("Nhân vật", level=1)
            for c in characters:
                p = doc.add_paragraph()
                name_run = p.add_run(c.name)
                name_run.bold = True
                role = getattr(c, "role", "") or ""
                if role:
                    role_run = p.add_run(f" ({role})")
                    role_run.italic = True
                personality = getattr(c, "personality", "") or ""
                if personality:
                    doc.add_paragraph(personality)
            doc.add_page_break()

        # Chapters
        for ch in story.chapters:
            heading_text = (
                f"Chương {ch.chapter_number}: {ch.title}"
                if getattr(ch, "chapter_number", None)
                else (ch.title or "")
            )
            doc.add_heading(heading_text, level=1)
            content = ch.content or ""
            for para in content.split("\n"):
                stripped = para.strip()
                if stripped:
                    doc.add_paragraph(stripped)
            doc.add_page_break()

        out_dir = os.path.dirname(output_path)
        if out_dir:
            os.makedirs(out_dir, exist_ok=True)
        doc.save(output_path)
        logger.info(f"DOCX exported: {output_path}")
        return output_path
